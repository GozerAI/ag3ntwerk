"""
Document Generation Integration for ag3ntwerk.

Provides document creation in multiple formats.

Requirements:
    - pip install python-docx markdown jinja2

Documents is ideal for:
    - Report generation
    - Contract creation
    - Template-based documents
    - Automated documentation
"""

import asyncio
import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional, Union
from enum import Enum
from pathlib import Path
import io

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """Document output formats."""

    DOCX = "docx"
    MARKDOWN = "md"
    HTML = "html"
    TXT = "txt"


@dataclass
class DocumentTemplate:
    """Represents a document template."""

    name: str
    content: str
    format: DocumentFormat = DocumentFormat.MARKDOWN
    variables: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TableData:
    """Represents table data for documents."""

    headers: List[str]
    rows: List[List[Any]]
    caption: str = ""


class DocumentGenerator:
    """
    Integration for document generation.

    Example:
        gen = DocumentGenerator()

        # Create from template
        doc = await gen.from_template(
            template="report.md",
            variables={"title": "Q4 Report", "data": metrics},
            output_format=DocumentFormat.DOCX,
        )

        # Create simple document
        await gen.create_document(
            output_path="report.docx",
            title="Monthly Report",
            sections=[
                {"heading": "Summary", "content": "..."},
                {"heading": "Data", "table": table_data},
            ],
        )
    """

    def __init__(self, templates_dir: Optional[str] = None):
        """Initialize document generator."""
        self.templates_dir = Path(templates_dir) if templates_dir else None
        self._templates: Dict[str, DocumentTemplate] = {}

    def register_template(self, template: DocumentTemplate) -> None:
        """Register a template."""
        self._templates[template.name] = template

    async def from_template(
        self,
        template_name: str,
        variables: Dict[str, Any],
        output_path: Optional[str] = None,
        output_format: DocumentFormat = DocumentFormat.DOCX,
    ) -> Union[str, bytes]:
        """
        Generate document from template.

        Args:
            template_name: Template name or path
            variables: Template variables
            output_path: Output file path
            output_format: Output format

        Returns:
            File path or document bytes
        """
        loop = asyncio.get_running_loop()

        def _generate():
            try:
                from jinja2 import Template, Environment, FileSystemLoader
            except ImportError:
                raise ImportError("jinja2 not installed. Install with: pip install jinja2")

            # Get template
            if template_name in self._templates:
                template_content = self._templates[template_name].content
            elif self.templates_dir and (self.templates_dir / template_name).exists():
                env = Environment(loader=FileSystemLoader(str(self.templates_dir)))
                template = env.get_template(template_name)
                rendered = template.render(**variables)
            else:
                # Assume template_name is the template content
                template = Template(template_name)
                rendered = template.render(**variables)

            # Convert to output format
            if output_format == DocumentFormat.MARKDOWN:
                result = rendered
            elif output_format == DocumentFormat.HTML:
                result = self._markdown_to_html(rendered)
            elif output_format == DocumentFormat.DOCX:
                result = self._markdown_to_docx(rendered)
            elif output_format == DocumentFormat.TXT:
                result = self._strip_markdown(rendered)
            else:
                result = rendered

            # Save or return
            if output_path:
                if isinstance(result, bytes):
                    with open(output_path, "wb") as f:
                        f.write(result)
                else:
                    with open(output_path, "w") as f:
                        f.write(result)
                return output_path
            return result

        return await loop.run_in_executor(None, _generate)

    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert Markdown to HTML."""
        try:
            import markdown

            return markdown.markdown(markdown_text, extensions=["tables", "fenced_code"])
        except ImportError:
            # Basic conversion
            html = markdown_text.replace("\n\n", "</p><p>")
            html = f"<p>{html}</p>"
            return html

    def _markdown_to_docx(self, markdown_text: str) -> bytes:
        """Convert Markdown to DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        doc = Document()

        lines = markdown_text.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                # Bullet list
                items = []
                while i < len(lines) and (lines[i].startswith("- ") or lines[i].startswith("* ")):
                    items.append(lines[i][2:])
                    i += 1
                i -= 1
                for item in items:
                    p = doc.add_paragraph(item, style="List Bullet")
            elif line.startswith("```"):
                # Code block
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif line.strip():
                doc.add_paragraph(line)

            i += 1

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _strip_markdown(self, markdown_text: str) -> str:
        """Strip Markdown formatting."""
        import re

        text = markdown_text

        # Remove headers
        text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)

        # Remove bold/italic
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"_(.+?)_", r"\1", text)

        # Remove links
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)

        # Remove code
        text = re.sub(r"`(.+?)`", r"\1", text)

        return text

    async def create_document(
        self,
        output_path: str,
        title: str,
        sections: List[Dict[str, Any]],
        author: str = "",
        date: Optional[datetime] = None,
        header: str = "",
        footer: str = "",
    ) -> str:
        """
        Create a document with sections.

        Args:
            output_path: Output file path
            title: Document title
            sections: List of section dicts with 'heading', 'content', 'table', 'list'
            author: Document author
            date: Document date
            header: Header text
            footer: Footer text

        Returns:
            Output file path
        """
        loop = asyncio.get_running_loop()
        output_format = DocumentFormat(Path(output_path).suffix.lstrip("."))

        def _create():
            if output_format == DocumentFormat.DOCX:
                return self._create_docx(output_path, title, sections, author, date, header, footer)
            else:
                return self._create_markdown(output_path, title, sections, author, date)

        return await loop.run_in_executor(None, _create)

    def _create_docx(
        self,
        output_path: str,
        title: str,
        sections: List[Dict[str, Any]],
        author: str,
        date: Optional[datetime],
        header: str,
        footer: str,
    ) -> str:
        """Create DOCX document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx not installed")

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        if author or date:
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if author:
                meta.add_run(f"Author: {author}")
            if date:
                if author:
                    meta.add_run(" | ")
                meta.add_run(f"Date: {date.strftime('%Y-%m-%d')}")

        doc.add_paragraph()  # Spacer

        # Sections
        for section in sections:
            if "heading" in section:
                doc.add_heading(section["heading"], level=1)

            if "content" in section:
                doc.add_paragraph(section["content"])

            if "table" in section:
                table_data = section["table"]
                if isinstance(table_data, TableData):
                    headers = table_data.headers
                    rows = table_data.rows
                else:
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])

                if headers and rows:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Table Grid"

                    # Headers
                    for i, header_text in enumerate(headers):
                        table.rows[0].cells[i].text = str(header_text)

                    # Rows
                    for row_data in rows:
                        row = table.add_row()
                        for i, cell_data in enumerate(row_data):
                            row.cells[i].text = str(cell_data)

            if "list" in section:
                for item in section["list"]:
                    doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph()  # Spacer between sections

        doc.save(output_path)
        return output_path

    def _create_markdown(
        self,
        output_path: str,
        title: str,
        sections: List[Dict[str, Any]],
        author: str,
        date: Optional[datetime],
    ) -> str:
        """Create Markdown document."""
        lines = []

        # Title
        lines.append(f"# {title}")
        lines.append("")

        # Metadata
        if author or date:
            meta_parts = []
            if author:
                meta_parts.append(f"*Author: {author}*")
            if date:
                meta_parts.append(f"*Date: {date.strftime('%Y-%m-%d')}*")
            lines.append(" | ".join(meta_parts))
            lines.append("")

        # Sections
        for section in sections:
            if "heading" in section:
                lines.append(f"## {section['heading']}")
                lines.append("")

            if "content" in section:
                lines.append(section["content"])
                lines.append("")

            if "table" in section:
                table_data = section["table"]
                if isinstance(table_data, TableData):
                    headers = table_data.headers
                    rows = table_data.rows
                else:
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])

                if headers and rows:
                    # Header row
                    lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                    # Separator
                    lines.append("| " + " | ".join("---" for _ in headers) + " |")
                    # Data rows
                    for row in rows:
                        lines.append("| " + " | ".join(str(c) for c in row) + " |")
                    lines.append("")

            if "list" in section:
                for item in section["list"]:
                    lines.append(f"- {item}")
                lines.append("")

        content = "\n".join(lines)

        with open(output_path, "w") as f:
            f.write(content)

        return output_path

    async def create_table_document(
        self,
        output_path: str,
        title: str,
        tables: List[TableData],
    ) -> str:
        """
        Create a document with multiple tables.

        Args:
            output_path: Output file path
            title: Document title
            tables: List of TableData

        Returns:
            Output file path
        """
        sections = []
        for table in tables:
            section = {"table": table}
            if table.caption:
                section["heading"] = table.caption
            sections.append(section)

        return await self.create_document(output_path, title, sections)
